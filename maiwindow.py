from os import path

__author__ = 'Killx'

import tkinter as tk
from tkinter import *
from tkinter.filedialog import askopenfilename
import webbrowser
import smtplib, base64
from win32com import client
import tkinter.messagebox as tkMessageBox

from docx import Document
from docx.shared import Pt
from docx.shared import Inches
from docx import *
from docx.enum.text import WD_LINE_SPACING
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_UNDERLINE

global mai, root, name, addres, wl2, emailid, phonee, weblink, univ, school, gpa, loc, course, major, datee, i, edu, adde, gen, nextedu, company, locexp, posi, startdate, enddate, detailss, exp, headln


def linkpage():
    webbrowser.open_new(r"http://resumebuild.herokuapp.com")


def updatepage():
    webbrowser.open_new(r"http://resumebuild.herokuapp.com")


def create_star():
    root.destroy()
    global mai
    mai = tk.Tk()
    mai.title("Create Resume")
    mai.geometry("1024x768")
    menubar = Menu(mai)
    mai.config(menu=menubar)

    fileMenu = Menu(menubar)
    menubar.add_cascade(label="File", menu=fileMenu)
    fileMenu.add_command(label="Open File        Ctrl+O")
    fileMenu.add_command(label="Save Form       Ctrl+S")
    fileMenu.add_command(label="Home Screen", command=root_app)
    fileMenu.add_command(label="Exit", command=mai.destroy)

    editMenu = Menu(menubar)
    menubar.add_cascade(label="Edit", menu=editMenu)
    editMenu.add_command(label="Copy", command="")
    editMenu.add_command(label="Paste", command="")
    editMenu.add_command(label="Select All", command="")

    aboutMenu = Menu(menubar)
    menubar.add_cascade(label="Help", menu=aboutMenu)
    aboutMenu.add_command(label="Help F1", command=create_faq1)
    aboutMenu.add_command(label="Check for updates", command=updatepage)
    aboutMenu.add_command(label="About", command=linkpage)

    global name, addres, wl2, emailid, phonee, weblink, univ, school, gpa, loc, course, major, datee, i, edu, adde, gen, nextedu, company, locexp, posi, startdate, enddate, detailss, exp, headln

    Label(text="Name").grid(row=1, column=0, padx=4, pady=4)
    name = Entry(mai, width=50)
    name.grid(column=1, columnspan=5, row=1, padx=4, pady=4)
    Label(text="Address").grid(row=1, column=6, padx=4, pady=4)
    addres = Entry(mai)
    addres.grid(column=7, columnspan=2, row=1, padx=4, pady=4)
    Label(text="Weblink").grid(row=1, column=9, padx=4, pady=4)
    wl2 = Entry(mai)
    wl2.grid(column=10, row=1, padx=4, pady=4)
    Label(text="Email ID").grid(row=2, column=0, padx=4, pady=4)
    emailid = Entry(mai, width=50)
    emailid.grid(column=1, columnspan=5, row=2, padx=4, pady=4)
    Label(text="Phone").grid(row=2, column=6, padx=4, pady=4)
    phonee = Entry(mai)
    phonee.grid(column=7, columnspan=2, row=2, padx=4, pady=4)
    Label(text="Web Link").grid(row=2, column=9, padx=4, pady=4)
    weblink = Entry(mai)
    weblink.grid(column=10, row=2, padx=4, pady=4)
    Label(text="Headline").grid(row=3, column=0)
    headln = Entry(mai, width=80)
    headln.grid(row=3, column=1, columnspan=8, padx=4, pady=4)
    Label(text="Education").grid(row=4, column=0)
    i = 4
    univ = [''] * 10
    school = [''] * 10
    gpa = [''] * 10
    loc = [''] * 10
    course = [''] * 10
    major = [''] * 10
    datee = [''] * 10
    edu = 0
    Label(text="University").grid(row=i, column=0)
    univ[edu] = Entry(mai)
    univ[edu].grid(row=i, column=1)
    Label(text="School").grid(row=i, column=2)
    school[edu] = Entry(mai)
    school[edu].grid(row=i, column=3)
    Label(text="GPA").grid(row=i, column=4)
    gpa[edu] = Entry(mai)
    gpa[edu].grid(row=i, column=5)
    Label(text="Location").grid(row=i, column=6)
    loc[edu] = Entry(mai)
    loc[edu].grid(row=i, column=7)
    i += 1
    Label(text="Course").grid(row=i, column=0)
    course[edu] = Entry(mai)
    course[edu].grid(row=i, column=1)
    Label(text="Major").grid(row=i, column=2)
    major[edu] = Entry(mai)
    major[edu].grid(row=i, column=3)
    Label(text="Duration").grid(row=i, column=4)
    datee[edu] = Entry(mai)
    datee[edu].grid(row=i, column=5)
    i += 1
    adde = Button(mai, text="Add Education", command=nexteduc)
    adde.grid(row=i, column=1, padx=4, pady=4)
    gen = Button(mai, text="Generate", command=gener)
    gen.grid(row=i, column=2, padx=4, pady=4)
    nextedu = Button(mai, text="Next", command=nextexp)
    nextedu.grid(row=i, column=4, padx=4, pady=4)

    mai.mainloop()


def nexteduc():
    global name, addres, wl2, emaiid, phonee, weblink, univ, school, gpa, loc, course, major, datee, i, edu, adde, gen, nextedu

    adde.destroy()
    gen.destroy()
    nextedu.destroy()
    edu += 1
    Label(text="University").grid(row=i, column=0)
    univ[edu] = Entry(mai)
    univ[edu].grid(row=i, column=1)
    Label(text="School").grid(row=i, column=2)
    school[edu] = Entry(mai)
    school[edu].grid(row=i, column=3)
    Label(text="GPA").grid(row=i, column=4)
    gpa[edu] = Entry(mai)
    gpa[edu].grid(row=i, column=5)
    Label(text="Location").grid(row=i, column=6)
    loc[edu] = Entry(mai)
    loc[edu].grid(row=i, column=7)
    i += 1
    Label(text="Course").grid(row=i, column=0)
    course[edu] = Entry(mai)
    course[edu].grid(row=i, column=1)
    Label(text="Major").grid(row=i, column=2)
    major[edu] = Entry(mai)
    major[edu].grid(row=i, column=3)
    Label(text="Duration").grid(row=i, column=4)
    datee[edu] = Entry(mai)
    datee[edu].grid(row=i, column=5)
    i += 1
    adde = Button(mai, text="Add Education", command=nexteduc)
    adde.grid(row=i, column=1, padx=4, pady=4)
    gen = Button(mai, text="Generate", command=gener)
    gen.grid(row=i, column=2, padx=4, pady=4)
    nextedu = Button(mai, text="Next", command=nextexp)
    nextedu.grid(row=i, column=4, padx=4, pady=4)


def nextexp():
    global i, exp, adde, gen, nextedu, company, locexp, posi, startdate, enddate, detailss

    adde.destroy()
    gen.destroy()
    nextedu.destroy()
    company = [''] * 7
    locexp = [''] * 7
    posi = [''] * 7
    startdate = [''] * 7
    enddate = [''] * 7
    detailss = [''] * 7
    Label(text="EXPERIENCE").grid(row=i, column=0)
    i += 1
    exp = 0
    Label(text="Company").grid(row=i, column=0)
    company[exp] = Entry(mai)
    company[exp].grid(row=i, column=1)
    Label(text="Location").grid(row=i, column=2)
    locexp[exp] = Entry(mai)
    locexp[exp].grid(row=i, column=3)
    Label(text="Position").grid(row=i, column=4)
    posi[exp] = Entry(mai)
    posi[exp].grid(row=i, column=5)
    i += 1
    Label(text="Start Date").grid(row=i, column=0)
    startdate[exp] = Entry(mai)
    startdate[exp].grid(row=i, column=1)
    Label(text="End Date").grid(row=i, column=2)
    enddate[exp] = Entry(mai)
    enddate[exp].grid(row=i, column=3)
    Label(text="Details").grid(row=i, column=4)
    detailss[exp] = Text(mai, width=60, height=2)
    detailss[exp].grid(row=i, column=5, columnspan=6)
    i += 1
    adde = Button(mai, text="Add Experience", command=nexteexp)
    adde.grid(row=i, column=1, padx=4, pady=4)
    gen = Button(mai, text="Generate", command=gener)
    gen.grid(row=i, column=2, padx=4, pady=4)
    nextedu = Button(mai, text="Next", command=nextskil)
    nextedu.grid(row=i, column=4, padx=4, pady=4)


def nexteexp():
    global i, exp, adde, gen, nextedu, company, locexp, posi, startdate, enddate, detailss

    adde.destroy()
    gen.destroy()
    nextedu.destroy()
    exp += 1
    Label(text="Company").grid(row=i, column=0)
    company[exp] = Entry(mai)
    company[exp].grid(row=i, column=1)
    Label(text="Location").grid(row=i, column=2)
    locexp[exp] = Entry(mai)
    locexp[exp].grid(row=i, column=3)
    Label(text="Position").grid(row=i, column=4)
    posi[exp] = Entry(mai)
    posi[exp].grid(row=i, column=5)
    i += 1
    Label(text="Start Date").grid(row=i, column=0)
    startdate[exp] = Entry(mai)
    startdate[exp].grid(row=i, column=1)
    Label(text="End Date").grid(row=i, column=2)
    enddate[exp] = Entry(mai)
    enddate[exp].grid(row=i, column=3)
    Label(text="Details").grid(row=i, column=4)
    detailss[exp] = Text(mai, width=60, height=2)
    detailss[exp].grid(row=i, column=5, columnspan=6)
    i += 1
    adde = Button(mai, text="Add Experience", command=nexteexp)
    adde.grid(row=i, column=1, padx=4, pady=4)
    gen = Button(mai, text="Generate", command=gener)
    gen.grid(row=i, column=2, padx=4, pady=4)
    nextedu = Button(mai, text="Next", command=nextskil)
    nextedu.grid(row=i, column=4, padx=4, pady=4)


def nextskil():
    global i, skil, stitle, sdetail, adde, gen, nextedu, projlab
    adde.destroy()
    gen.destroy()
    nextedu.destroy()
    stitle = [''] * 5
    sdetail = [''] * 5
    skil = 0
    Label(text="Skills").grid(row=i, column=0)
    i += 1
    Label(text="Title").grid(row=i, column=0)
    stitle[skil] = Entry(mai)
    stitle[skil].grid(row=i, column=1)
    Label(text="Details").grid(row=i, column=2)
    sdetail[skil] = Entry(mai, width=90)
    sdetail[skil].grid(row=i, column=3, columnspan=7)
    i += 1
    adde = Button(mai, text="Add Skill Set", command=nextskill)
    adde.grid(row=i, column=1, padx=4, pady=4)
    projlab = Button(mai, text=" Next Project", command=proj)
    projlab.grid(row=i, column=2, padx=4, pady=4)


def nextskill():
    global i, skil, stitle, sdetail, adde, gen, nextedu, projlab
    adde.destroy()
    gen.destroy()
    projlab.destroy()

    skil += 1

    Label(text="Title").grid(row=i, column=0)
    stitle[skil] = Entry(mai)
    stitle[skil].grid(row=i, column=1)
    Label(text="Details").grid(row=i, column=2)
    sdetail[skil] = Entry(mai, width=90)
    sdetail[skil].grid(row=i, column=3, columnspan=7)
    i += 1
    adde = Button(mai, text="Add Skill Set", command=nextskill)
    adde.grid(row=i, column=1, padx=4, pady=4)
    projlab = Button(mai, text=" Next Project", command=proj)
    projlab.grid(row=i, column=2, padx=4, pady=4)


def proj():
    global i, projcount, projname, projdetail, adde, gen, genword, genhtml, genpdf, addproj, projlab
    adde.destroy()
    projlab.destroy()
    gen.destroy()
    projname = [''] * 5
    projdetail = [''] * 5
    projcount = 0
    Label(text="Projects").grid(row=i, column=0)
    i += 1
    Label(text="Name").grid(row=i, column=0)
    projname[projcount] = Entry(mai)
    projname[projcount].grid(row=i, column=1)
    Label(text="Details").grid(row=i, column=2)
    projdetail[projcount] = Entry(mai, width=90)
    projdetail[projcount].grid(row=i, column=3, columnspan=7)
    i += 1
    addproj = Button(mai, text="Add Project", command=nextproj)
    addproj.grid(row=i, column=1, padx=4, pady=4)
    gen = Button(mai, text="Generate All", command=gener)
    gen.grid(row=i, column=2, padx=4, pady=4)
    genword = Button(mai, text="Generate Word", command=generword)
    genword.grid(row=i, column=3, padx=4, pady=4)
    genhtml = Button(mai, text="Generate Word", command=generhtml)
    genhtml.grid(row=i, column=4, padx=4, pady=4)
    genpdf = Button(mai, text="Generate Word", command=generpdf)
    genpdf.grid(row=i, column=5, padx=4, pady=4)


def nextproj():
    global i, projcount, projname, projdetail, adde, gen, genword, genhtml, genpdf, addproj
    gen.destroy()
    genpdf.destroy()
    genhtml.destroy()
    genword.destroy()
    addproj.destroy()
    projcount += 1
    i += 1
    Label(text="Name").grid(row=i, column=0)
    projname[projcount] = Entry(mai)
    projname[projcount].grid(row=i, column=1)
    Label(text="Details").grid(row=i, column=2)
    projdetail[projcount] = Entry(mai, width=90)
    projdetail[projcount].grid(row=i, column=3, columnspan=7)
    i += 1
    addproj = Button(mai, text="Add Project", command=nextproj)
    addproj.grid(row=i, column=1, padx=4, pady=4)
    gen = Button(mai, text="Generate All", command=gener)
    gen.grid(row=i, column=2, padx=4, pady=4)
    genword = Button(mai, text="Generate Word", command=generword)
    genword.grid(row=i, column=3, padx=4, pady=4)
    genhtml = Button(mai, text="Generate Word", command=generhtml)
    genhtml.grid(row=i, column=4, padx=4, pady=4)
    genpdf = Button(mai, text="Generate Word", command=generpdf)
    genpdf.grid(row=i, column=5, padx=4, pady=4)


def generword():
    gname = name.get()
    gaddres = addres.get()
    gwl2 = wl2.get()
    gemail = emailid.get()
    gphonee = phonee.get()
    gweblink = weblink.get()
    gheadln = headln.get()
    #word_file

    document = Document()

    # name
    head = document.add_paragraph()
    head.add_run(gname)
    head.paragraph_format.space_before = Pt(0)
    head.paragraph_format.space_after = Pt(0)
    head.size = Pt(14)
    head.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    #address and phone
    paragrapha = document.add_paragraph(gaddres + gphonee)

    paragrapha.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    paragrapha.paragraph_format.space_before = Pt(0)
    paragrapha.paragraph_format.space_after = Pt(0)
    paragrapha.size = Pt(12)

    #email and weblinks
    paragrapha = document.add_paragraph(gweblink + gemail)

    paragrapha.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    paragrapha.paragraph_format.space_before = Pt(0)
    paragrapha.paragraph_format.space_after = Pt(8)
    paragrapha.size = Pt(12)


    edunmaee = document.add_paragraph()
    edunmaee.size = Pt(12)
    edunmaee.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    edunmaee.paragraph_format.space_before = Pt(0)
    edunmaee.paragraph_format.space_after = Pt(0)
    edunmaee = paragrapha.add_run('\n EDUCATION')
    edunmaee.bold = True
    edunmaee.underline = WD_UNDERLINE.SINGLE
    if edu > 0:
        for ed in range(0, edu + 1):
            guniv = univ[ed].get()
            gschool = school[ed].get()
            ggpa = gpa[ed].get()
            gloc = loc[ed].get()
            gcourse = course[ed].get()
            gmajor = major[ed].get()
            gdate = datee[ed].get()
            paragrapha = document.add_paragraph()
            paragrapha.paragraph_format.space_before = Pt(0)
            paragrapha.paragraph_format.space_after = Pt(0)
            paragrapha.size = Pt(12)
            paragrapha.add_run(guniv + "," + gschool + ",").bold = True
            paragrapha.add_run(gloc)

            paragrapha = document.add_paragraph()
            paragrapha.paragraph_format.space_before = Pt(0)
            paragrapha.paragraph_format.space_after = Pt(0)
            paragrapha.size = Pt(12)
            paragrapha.add_run(gcourse + "," + gmajor + " | " + ggpa + "     " + gdate)

    elif edu == 0:
        guniv = univ[0].get()
        gschool = school[0].get()
        ggpa = gpa[0].get()
        gloc = loc[0].get()
        gcourse = course[0].get()
        gmajor = major[0].get()
        gdate = datee[0].get()
        paragrapha = document.add_paragraph()
        paragrapha.paragraph_format.space_before = Pt(0)
        paragrapha.paragraph_format.space_after = Pt(0)
        paragrapha.size = Pt(12)
        paragrapha.add_run(guniv + "," + gschool + ",").bold = True
        paragrapha.add_run(gloc)

        paragrapha = document.add_paragraph()
        paragrapha.paragraph_format.space_before = Pt(0)
        paragrapha.paragraph_format.space_after = Pt(0)
        paragrapha.size = Pt(12)
        paragrapha.add_run(gcourse + "," + gmajor + " | " + ggpa + "     " + gdate)





    edunmaee = document.add_paragraph()
    edunmaee.size = Pt(12)
    edunmaee.paragraph_format.space_before = Pt(0)
    edunmaee.paragraph_format.space_after = Pt(0)
    edunmaee = paragrapha.add_run('EXPERIENCE')
    edunmaee.bold = True
    edunmaee.underline = WD_UNDERLINE.SINGLE
    if exp > 0:
        for ed in range(0, exp + 1):
            gcomp = company[ed].get()
            glocexp = locexp[ed].get()
            gposi = posi[ed].get()
            gsd = startdate[ed].get()
            ged = enddate[ed].get()
            gdetails = detailss[ed].get(1.0, END)
            paragrapha = document.add_paragraph()
            paragrapha.paragraph_format.space_before = Pt(0)
            paragrapha.paragraph_format.space_after = Pt(0)
            paragrapha.size = Pt(12)
            paragrapha.add_run(gcomp + "," + glocexp).bold = True

            paragrapha = document.add_paragraph()
            paragrapha.paragraph_format.space_before = Pt(0)
            paragrapha.paragraph_format.space_after = Pt(0)
            paragrapha.size = Pt(12)
            paragrapha.underline = WD_UNDERLINE.SINGLE
            paragrapha.add_run(gposi)
            paragrapha.underline = WD_UNDERLINE.NONE
            paragrapha.add_run(", "+ gsd + " - " + ged)

            paragrapha = document.add_paragraph()
            paragrapha.paragraph_format.space_before = Pt(0)
            paragrapha.paragraph_format.space_after = Pt(0)
            paragrapha.size = Pt(12)
            paragrapha.add_run(gdetails)


    elif exp == 0:
            gcomp = company[0].get()
            glocexp = locexp[0].get()
            gposi = posi[0].get()
            gsd = startdate[0].get()
            ged = enddate[0].get()
            gdetails = detailss[0].get(1.0,END)
            paragrapha = document.add_paragraph()
            paragrapha.paragraph_format.space_before = Pt(0)
            paragrapha.paragraph_format.space_after = Pt(0)
            paragrapha.size = Pt(12)
            paragrapha.add_run(gcomp + "," + glocexp).bold = True

            paragrapha = document.add_paragraph()
            paragrapha.paragraph_format.space_before = Pt(0)
            paragrapha.paragraph_format.space_after = Pt(0)
            paragrapha.size = Pt(12)
            paragrapha.underline = WD_UNDERLINE.SINGLE
            paragrapha.add_run(gposi)
            paragrapha.underline = WD_UNDERLINE.NONE
            paragrapha.add_run(", "+ gsd + " - " + ged)
            paragrapha = document.add_paragraph()
            paragrapha.paragraph_format.space_before = Pt(0)
            paragrapha.paragraph_format.space_after = Pt(0)
            paragrapha.size = Pt(12)
            paragrapha.add_run(gdetails)



    edunmaee = document.add_paragraph()
    edunmaee.size = Pt(12)
    paragrapha.paragraph_format.space_before = Pt(2)
    paragrapha.paragraph_format.space_after = Pt(4)
    edunmaee = paragrapha.add_run('SKILLS')
    edunmaee.bold = True
    edunmaee.underline = WD_UNDERLINE.SINGLE
    if skil > 0:
        for sksk in range(0, skil + 1):
            gstitle = stitle[sksk].get()
            gsdetail = sdetail[sksk].get()
            paragrapha = document.add_paragraph()
            paragrapha.paragraph_format.space_before = Pt(0)
            paragrapha.paragraph_format.space_after = Pt(0)
            paragrapha.size = Pt(12)
            paragrapha.add_run(gstitle).bold = True
            paragrapha.add_run(" : " + gsdetail).bold = False


    elif skil == 0:
        gstitle = stitle[0].get()
        gsdetail = sdetail[0].get()
        paragrapha = document.add_paragraph()
        paragrapha.paragraph_format.space_before = Pt(0)
        paragrapha.paragraph_format.space_after = Pt(0)
        paragrapha.size = Pt(12)
        paragrapha.add_run(gstitle).bold = True
        paragrapha.add_run(" : " + gsdetail).bold = False


    edunmaee = document.add_paragraph()
    edunmaee.size = Pt(12)
    edunmaee.paragraph_format.space_before = Pt(2)
    edunmaee.paragraph_format.space_after = Pt(4)
    edunmaee = paragrapha.add_run('PROJECTS')
    edunmaee.bold = True
    edunmaee.underline = WD_UNDERLINE.SINGLE
    if projcount > 0:
        for pjpj in range(0, projcount + 1):
            gprojtitle = projname[pjpj].get()
            gprojdetail = projdetail[pjpj].get()
            paragrapha = document.add_paragraph()
            paragrapha.paragraph_format.space_before = Pt(0)
            paragrapha.paragraph_format.space_after = Pt(0)
            paragrapha.size = Pt(12)
            paragrapha.add_run(gprojtitle).bold = True
            paragrapha.add_run(" : " + gprojdetail).bold = False


    elif projcount == 0:
        gprojtitle = projname[0].get()
        gprojdetail = projdetail[0].get()
        paragrapha = document.add_paragraph()
        paragrapha.paragraph_format.space_before = Pt(0)
        paragrapha.paragraph_format.space_after = Pt(0)
        paragrapha.size = Pt(12)
        paragrapha.add_run(gprojtitle).bold = True
        paragrapha.add_run(" : " + gprojdetail).bold = False







    document.save('demo.docx')
    completeword()


def generhtml():
    global name, addres, headln, wl2, emailid, phonee, weblink, univ, school, gpa, loc, course, major, datee, exp, edu, adde, gen, nextedu, company, locexp, posi, startdate, enddate, detailss, skil, stitle, sdetail, projcount, projname, projdetail
    gname = name.get()
    gaddres = addres.get()
    gwl2 = wl2.get()
    gemail = emailid.get()
    gphonee = phonee.get()
    gweblink = weblink.get()
    gheadln = headln.get()

    htmlfile = open("resume.html", "w")
    html1 = """
    <!DOCTYPE html>

    <html>

    <head>
         <meta http-equiv="Content-Type" content="text/html; charset=utf-8"/>
    """
    htmlfile.write(html1)
    html2 = ("<title>" + gname + "</title>")
    htmlfile.write(html2)
    html3 = """
        <style type="text/css">
            * { margin: 0; padding: 0; }
            body { font: 16px Helvetica, Sans-Serif; line-height: 24px; background: url(images/noise.jpg); }
            .clear { clear: both; }
            .vcard {padding-left: 9px;}

            #pic { float: right; margin: -30px 0 0 0;width: 15%;min-width:100px;height: auto;padding-right: 25px; padding-bottom:5px;opacity: .7;border-radius: 10px;border: 3px solid white; }
            h1 { margin: 0 0 16px 0; padding: 0 0 16px 0; font-size: 42px; font-weight: bold; letter-spacing: -2px; }
            h2 { font-size: 20px; margin: 0 0 6px 0; position: relative; }
            h2 span { position: relative; bottom: 0; right: 0; font-style: italic; font-family: Georgia, Serif; font-size: 16px; color: #999; font-weight: normal; }
            p { margin: 0 0 16px 0; }
            a { color: #999; text-decoration: none; border-bottom: 1px dotted #999; }
            a:hover { border-bottom-style: solid; color: black; }
            ul { margin: 0 0 32px 17px; }

            #objective p { font-family: Georgia, Serif; font-style: italic; color: #666; }

            ul {padding-left: 10px;}
            footer{text-align: center;margin-bottom: 20px;}
            footer h1 {font-size: 20px;margin-top: 0px;margin-bottom: 0px}
            footer ul{list-style:none;margin-bottom: 10px;margin-top: 0px;}
            footer ul li {display: inline;}
            #soc {width:48px; height:48px;margin-left:10px;margin-right: 10px;-webkit-transition: all 1s ease; -webkit-filter: grayscale(100%); }
            #soc:hover {-webkit-filter: grayscale(0%);
                -webkit-transform: rotate(-10deg);}

            @media screen and (max-width: 768px){
            #objective { width: 85%; float: left; }
            dt { font-style: italic; font-weight: bold; font-size: 18px; text-align: right; padding: 0 26px 0 0;width:90%; margin-left: 3% ;border-bottom: 2px solid black;margin-bottom: 10px;text-align: center;}
            dd { padding-left: 26px;width:75%;}
            dd.clear { float: none; margin: 0; height: 45px;border: none; }
            #page-wrap {width:90%;margin: 20px;}

            }

            @media screen and (min-width: 769px){
                 dt { font-style: italic; font-weight: bold; font-size: 18px; text-align: right; padding: 0 26px 0 0;width:150px;float: left; margin-left: 10% }
            dd { float: left;border-left: 1px solid #999; padding-left: 26px;width:60%;}
            dd.clear { float: none; margin: 0; height: 15px;border: none; }
             #objective { width: 55%; float: left; }
             #page-wrap { width: 80%; margin: 40px auto 60px; }

            }

         </style>
    </head>

    <body>

        <div id="page-wrap">



            <div id="contact-info" class="vcard">
            <img src="img.jpg" alt="Photo of D" id="pic" />

                <!-- Microformats! -->

                <h1 class="fn">
    """
    htmlfile.write(html3)
    html4 = """
    </h1>

            <p>
                Cell: <span class="tel">
    """
    html5 = """
    </span><br />
                Email: <a class="email" href="mailto:
    """
    html6 = """
    ">
    """
    html7 = """
    </a>
            </p>
        </div>

        <div id="objective">
            <p>
    """
    htmlfile.write(gname + html4 + gphonee + html5 + gemail + html6 + gemail + html7)
    html8 = """
     </p>
        </div>
        <hr style="clear:both;">
        <div class="clear"></div>

        <dl>
            <dd class="clear"></dd>

            <dt>Education</dt><dd>

    """
    htmlfile.write(gheadln + html8)

    if edu > 0:
        for ed in range(0, edu + 1):
            guniv = univ[ed].get()
            gschool = school[ed].get()
            ggpa = gpa[ed].get()
            gloc = loc[ed].get()
            gcourse = course[ed].get()
            gmajor = major[ed].get()
            gdate = datee[ed].get()
            html10 = """
            <h2>
            """
            html11 = ","
            html12 = "</h2><p><strong>"
            html13 = "</strong>"
            html14 = "<br /><strong>Major:</strong>"
            html15 = "</p>"
            htmlfile.write(
                html10 + guniv + html11 + gschool + html11 + gloc + html12 + gcourse + html13 + gdate + html14 + gmajor + html15)


    elif edu == 0:
        guniv = univ[0].get()
        gschool = school[0].get()
        ggpa = gpa[0].get()
        gloc = loc[0].get()
        gcourse = course[0].get()
        gmajor = major[0].get()
        gdate = datee[0].get()
        html10 = """
            <h2>
            """
        html11 = ","
        html12 = "</h2><p><strong>"
        html13 = "</strong>"
        html14 = "<br /><strong>Major:</strong>"
        html15 = "</p></dd>"
        htmlfile.write(html10)
        htmlfile.write(guniv)
        htmlfile.write(html11)
        htmlfile.write(gschool)
        htmlfile.write(html11)
        htmlfile.write(gloc)
        htmlfile.write(html12)
        htmlfile.write(gcourse)
        htmlfile.write(html13)
        htmlfile.write(gdate)
        htmlfile.write(html14)
        htmlfile.write(gmajor)
        htmlfile.write(html15)
    html16 = "</dd>"
    htmlfile.write(html16)
    html17 = """<dd class="clear"></dd>
    <dt>Experience</dt>
    <dd>
    """
    htmlfile.write(html17)

    if exp > 0:
        for ed in range(0, exp + 1):
            gcomp = company[ed].get()
            glocexp = locexp[ed].get()
            gposi = posi[ed].get()
            gsd = startdate[ed].get()
            ged = enddate[ed].get()
            gdetails = detailss[ed].get(1.0, END)
            html10 = "<h2>"
            html11 = "<br><span>"
            html12 = " - "
            html13 = """</span><br></h2>

                <ul>
                    <li>"""
            html14 = "</li></ul>"

            htmlfile.write(
                html10 + gcomp + html11 + gposi + html12 + glocexp + html12 + gsd + html12 + ged + html13 + gdetails + html14)


    elif exp == 0:
        gcomp = company[0].get()
        glocexp = locexp[0].get()
        gposi = posi[0].get()
        gsd = startdate[0].get()
        ged = enddate[0].get()
        gdetails = detailss[0].get(1.0, END)
        html10 = "<h2>"
        html11 = "<br><span>"
        html12 = " - "
        html13 = """</span><br></h2>
            <ul>
                    <li>"""
        html14 = "</li></ul>"

        htmlfile.write(
            html10 + gcomp + html11 + gposi + html12 + glocexp + html12 + gsd + html12 + ged + html13 + gdetails + html14)

    htmlfile.write(html16)
    html17 = """<dd class="clear"></dd>
    <dt>Skills</dt>
    <dd>
    """
    htmlfile.write(html17)

    if skil > 0:
        for sksk in range(0, skil + 1):
            gstitle = stitle[sksk].get()
            gsdetail = sdetail[sksk].get()
            html10 = "<h2>"
            html11 = "</h2><p>"
            html12 = "</p>"
            htmlfile.write(
                html10 + gstitle + html11 + gsdetail + html12)

    elif skil == 0:
        gstitle = stitle[0].get()
        gsdetail = sdetail[0].get()
        html10 = "<h2>"
        html11 = "</h2><p>"
        html12 = "</p>"
        htmlfile.write(
            html10 + gstitle + html11 + gsdetail + html12)

    htmlfile.write(html16)
    html17 = """<dd class="clear"></dd>
    <dt>Projects</dt>
    <dd>
    """
    htmlfile.write(html17)

    if projcount > 0:
        for pjpj in range(0, projcount + 1):
            gprojtitle = projname[pjpj].get()
            gprojdetail = projdetail[pjpj].get()
            html10 = "<h2>"
            html11 = "</h2><p>"
            html12 = "</p>"
            htmlfile.write(
                html10 + gprojtitle + html11 + gprojdetail + html12)

    elif projcount == 0:
        gprojtitle = projname[0].get()
        gprojdetail = projdetail[0].get()
        html10 = "<h2>"
        html11 = "</h2><p>"
        html12 = "</p>"
        htmlfile.write(
            html10 + gprojtitle + html11 + gprojdetail + html12)

    httml = """
        <dd class="clear"></dd>



                <dt>References</dt>
                <dd>Available on request</dd>




            </dl>
            <div class="clear"></div>
            <hr style="margin-bottom:0px;margin-top:25px">


        </div>
        <footer>
            <h1><a href="
        """
    htmlfile.write(httml + gwl2)
    httml = """
    ">
    """
    httml1 = "</h1><br></footer></body></html>"
    htmlfile.write(httml + gname + httml1)
    htmlfile.close()
    completehtml()


def generpdf():
    global name, addres, headln, wl2, emailid, phonee, weblink, univ, school, gpa, loc, course, major, datee, exp, edu, adde, gen, nextedu, company, locexp, posi, startdate, enddate, detailss, skil, stitle, sdetail, projcount, projname, projdetail
    gname = name.get()
    gaddres = addres.get()
    gwl2 = wl2.get()
    gemail = emailid.get()
    gphonee = phonee.get()
    gweblink = weblink.get()
    gheadln = headln.get()
    #word_file

    document = Document()

    # name
    head = document.add_paragraph()
    head.add_run(gname)
    head.paragraph_format.space_before = Pt(0)
    head.paragraph_format.space_after = Pt(0)
    head.size = Pt(14)
    head.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    #address and phone
    paragrapha = document.add_paragraph(gaddres + gphonee)

    paragrapha.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    paragrapha.paragraph_format.space_before = Pt(0)
    paragrapha.paragraph_format.space_after = Pt(0)
    paragrapha.size = Pt(12)

    #email and weblinks
    paragrapha = document.add_paragraph(gweblink + gemail)

    paragrapha.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    paragrapha.paragraph_format.space_before = Pt(0)
    paragrapha.paragraph_format.space_after = Pt(8)
    paragrapha.size = Pt(12)


    edunmaee = document.add_paragraph()
    edunmaee.size = Pt(12)
    edunmaee.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    edunmaee.paragraph_format.space_before = Pt(0)
    edunmaee.paragraph_format.space_after = Pt(0)
    edunmaee.add_run('EDUCATION')
    edunmaee.bold = True
    edunmaee.underline = WD_UNDERLINE.SINGLE
    if edu > 0:
        for ed in range(0, edu + 1):
            guniv = univ[ed].get()
            gschool = school[ed].get()
            ggpa = gpa[ed].get()
            gloc = loc[ed].get()
            gcourse = course[ed].get()
            gmajor = major[ed].get()
            gdate = datee[ed].get()
            paragrapha = document.add_paragraph()
            paragrapha.paragraph_format.space_before = Pt(0)
            paragrapha.paragraph_format.space_after = Pt(0)
            paragrapha.size = Pt(12)
            paragrapha.add_run(guniv + "," + gschool + ",").bold = True
            paragrapha.add_run(gloc)

            paragrapha = document.add_paragraph()
            paragrapha.paragraph_format.space_before = Pt(0)
            paragrapha.paragraph_format.space_after = Pt(0)
            paragrapha.size = Pt(12)
            paragrapha.add_run(gcourse + "," + gmajor + " | " + ggpa + "     " + gdate)

    elif edu == 0:
        guniv = univ[0].get()
        gschool = school[0].get()
        ggpa = gpa[0].get()
        gloc = loc[0].get()
        gcourse = course[0].get()
        gmajor = major[0].get()
        gdate = datee[0].get()
        paragrapha = document.add_paragraph()
        paragrapha.paragraph_format.space_before = Pt(0)
        paragrapha.paragraph_format.space_after = Pt(0)
        paragrapha.size = Pt(12)
        paragrapha.add_run(guniv + "," + gschool + ",").bold = True
        paragrapha.add_run(gloc)

        paragrapha = document.add_paragraph()
        paragrapha.paragraph_format.space_before = Pt(0)
        paragrapha.paragraph_format.space_after = Pt(0)
        paragrapha.size = Pt(12)
        paragrapha.add_run(gcourse + "," + gmajor + " | " + ggpa + "     " + gdate)





    edunmaee = document.add_paragraph()
    edunmaee.size = Pt(12)
    edunmaee.paragraph_format.space_before = Pt(0)
    edunmaee.paragraph_format.space_after = Pt(0)
    edunmaee.add_run('EXPERIENCE')
    edunmaee.bold = True
    edunmaee.underline = WD_UNDERLINE.SINGLE
    if exp > 0:
        for ed in range(0, exp + 1):
            gcomp = company[ed].get()
            glocexp = locexp[ed].get()
            gposi = posi[ed].get()
            gsd = startdate[ed].get()
            ged = enddate[ed].get()
            gdetails = detailss[ed].get(1.0, END)
            paragrapha = document.add_paragraph()
            paragrapha.paragraph_format.space_before = Pt(0)
            paragrapha.paragraph_format.space_after = Pt(0)
            paragrapha.size = Pt(12)
            paragrapha.add_run(gcomp + "," + glocexp).bold = True

            paragrapha = document.add_paragraph()
            paragrapha.paragraph_format.space_before = Pt(0)
            paragrapha.paragraph_format.space_after = Pt(0)
            paragrapha.size = Pt(12)
            paragrapha.underline = WD_UNDERLINE.SINGLE
            paragrapha.add_run(gposi)
            paragrapha.underline = WD_UNDERLINE.NONE
            paragrapha.add_run(", "+ gsd + " - " + ged)

            paragrapha = document.add_paragraph()
            paragrapha.paragraph_format.space_before = Pt(0)
            paragrapha.paragraph_format.space_after = Pt(0)
            paragrapha.size = Pt(12)
            paragrapha.add_run(gdetails)


    elif exp == 0:
            gcomp = company[0].get()
            glocexp = locexp[0].get()
            gposi = posi[0].get()
            gsd = startdate[0].get()
            ged = enddate[0].get()
            gdetails = detailss[0].get(1.0,END)
            paragrapha = document.add_paragraph()
            paragrapha.paragraph_format.space_before = Pt(0)
            paragrapha.paragraph_format.space_after = Pt(0)
            paragrapha.size = Pt(12)
            paragrapha.add_run(gcomp + "," + glocexp).bold = True

            paragrapha = document.add_paragraph()
            paragrapha.paragraph_format.space_before = Pt(0)
            paragrapha.paragraph_format.space_after = Pt(0)
            paragrapha.size = Pt(12)
            paragrapha.underline = WD_UNDERLINE.SINGLE
            paragrapha.add_run(gposi)
            paragrapha.underline = WD_UNDERLINE.NONE
            paragrapha.add_run(", "+ gsd + " - " + ged)
            paragrapha = document.add_paragraph()
            paragrapha.paragraph_format.space_before = Pt(0)
            paragrapha.paragraph_format.space_after = Pt(0)
            paragrapha.size = Pt(12)
            paragrapha.add_run(gdetails)



    edunmaee = document.add_paragraph()
    edunmaee.size = Pt(12)
    edunmaee.paragraph_format.space_before = Pt(2)
    edunmaee.paragraph_format.space_after = Pt(4)
    edunmaee.add_run('SKILLS')
    edunmaee.bold = True
    edunmaee.underline = WD_UNDERLINE.SINGLE
    if skil > 0:
        for sksk in range(0, skil + 1):
            gstitle = stitle[sksk].get()
            gsdetail = sdetail[sksk].get()
            paragraphab = document.add_paragraph()
            paragraphab.paragraph_format.space_before = Pt(0)
            paragraphab.paragraph_format.space_after = Pt(0)
            paragraphab.size = Pt(12)
            paragraphab.add_run(gstitle).bold = True
            paragraphab.add_run(" : " + gsdetail).bold = False


    elif skil == 0:
        gstitle = stitle[0].get()
        gsdetail = sdetail[0].get()
        paragrapha = document.add_paragraph()
        paragrapha.paragraph_format.space_before = Pt(0)
        paragrapha.paragraph_format.space_after = Pt(0)
        paragrapha.size = Pt(12)
        paragrapha.add_run(gstitle).bold = True
        paragrapha.add_run(" : " + gsdetail).bold = False


    edunmaee = document.add_paragraph()
    edunmaee.size = Pt(12)
    edunmaee.paragraph_format.space_before = Pt(0)
    edunmaee.paragraph_format.space_after = Pt(0)
    edunmaee.bold = True
    edunmaee.underline = WD_UNDERLINE.SINGLE
    edunmaee.add_run('PROJECTS')

    if projcount > 0:
        for pjpj in range(0, projcount + 1):
            gprojtitle = projname[pjpj].get()
            gprojdetail = projdetail[pjpj].get()
            paragrapha = document.add_paragraph()
            paragrapha.paragraph_format.space_before = Pt(0)
            paragrapha.paragraph_format.space_after = Pt(0)
            paragrapha.size = Pt(12)
            paragrapha.add_run(gprojtitle).bold = True
            paragrapha.add_run(" : " + gprojdetail).bold = False


    elif projcount == 0:
        gprojtitle = projname[0].get()
        gprojdetail = projdetail[0].get()
        paragrapha = document.add_paragraph()
        paragrapha.paragraph_format.space_before = Pt(0)
        paragrapha.paragraph_format.space_after = Pt(0)
        paragrapha.size = Pt(12)
        paragrapha.add_run(gprojtitle).bold = True
        paragrapha.add_run(" : " + gprojdetail).bold = False







    document.save('demo.docx')
    completepdf()
    pdffi = path.abspath('.\demo.docx')
    convert_to_pdf(pdffi)


def gener():
    global name, addres, headln, wl2, emailid, phonee, weblink, univ, school, gpa, loc, course, major, datee, exp, edu, adde, gen, nextedu, company, locexp, posi, startdate, enddate, detailss, skil, stitle, sdetail, projcount, projname, projdetail
    gname = name.get()
    gaddres = addres.get()
    gwl2 = wl2.get()
    gemail = emailid.get()
    gphonee = phonee.get()
    gweblink = weblink.get()
    gheadln = headln.get()

    htmlfile = open("resume.html", "w")
    html1 = """
    <!DOCTYPE html>

    <html>

    <head>
         <meta http-equiv="Content-Type" content="text/html; charset=utf-8"/>
    """
    htmlfile.write(html1)
    html2 = ("<title>" + gname + "</title>")
    htmlfile.write(html2)
    html3 = """
        <style type="text/css">
            * { margin: 0; padding: 0; }
            body { font: 16px Helvetica, Sans-Serif; line-height: 24px; background: url(images/noise.jpg); }
            .clear { clear: both; }
            .vcard {padding-left: 9px;}

            #pic { float: right; margin: -30px 0 0 0;width: 15%;min-width:100px;height: auto;padding-right: 25px; padding-bottom:5px;opacity: .7;border-radius: 10px;border: 3px solid white; }
            h1 { margin: 0 0 16px 0; padding: 0 0 16px 0; font-size: 42px; font-weight: bold; letter-spacing: -2px; }
            h2 { font-size: 20px; margin: 0 0 6px 0; position: relative; }
            h2 span { position: relative; bottom: 0; right: 0; font-style: italic; font-family: Georgia, Serif; font-size: 16px; color: #999; font-weight: normal; }
            p { margin: 0 0 16px 0; }
            a { color: #999; text-decoration: none; border-bottom: 1px dotted #999; }
            a:hover { border-bottom-style: solid; color: black; }
            ul { margin: 0 0 32px 17px; }

            #objective p { font-family: Georgia, Serif; font-style: italic; color: #666; }

            ul {padding-left: 10px;}
            footer{text-align: center;margin-bottom: 20px;}
            footer h1 {font-size: 20px;margin-top: 0px;margin-bottom: 0px}
            footer ul{list-style:none;margin-bottom: 10px;margin-top: 0px;}
            footer ul li {display: inline;}
            #soc {width:48px; height:48px;margin-left:10px;margin-right: 10px;-webkit-transition: all 1s ease; -webkit-filter: grayscale(100%); }
            #soc:hover {-webkit-filter: grayscale(0%);
                -webkit-transform: rotate(-10deg);}

            @media screen and (max-width: 768px){
            #objective { width: 85%; float: left; }
            dt { font-style: italic; font-weight: bold; font-size: 18px; text-align: right; padding: 0 26px 0 0;width:90%; margin-left: 3% ;border-bottom: 2px solid black;margin-bottom: 10px;text-align: center;}
            dd { padding-left: 26px;width:75%;}
            dd.clear { float: none; margin: 0; height: 45px;border: none; }
            #page-wrap {width:90%;margin: 20px;}

            }

            @media screen and (min-width: 769px){
                 dt { font-style: italic; font-weight: bold; font-size: 18px; text-align: right; padding: 0 26px 0 0;width:150px;float: left; margin-left: 10% }
            dd { float: left;border-left: 1px solid #999; padding-left: 26px;width:60%;}
            dd.clear { float: none; margin: 0; height: 15px;border: none; }
             #objective { width: 55%; float: left; }
             #page-wrap { width: 80%; margin: 40px auto 60px; }

            }

         </style>
    </head>

    <body>

        <div id="page-wrap">



            <div id="contact-info" class="vcard">
            <img src="img.jpg" alt="Photo of D" id="pic" />

                <!-- Microformats! -->

                <h1 class="fn">
    """
    htmlfile.write(html3)
    html4 = """
    </h1>

            <p>
                Cell: <span class="tel">
    """
    html5 = """
    </span><br />
                Email: <a class="email" href="mailto:
    """
    html6 = """
    ">
    """
    html7 = """
    </a>
            </p>
        </div>

        <div id="objective">
            <p>
    """
    htmlfile.write(gname + html4 + gphonee + html5 + gemail + html6 + gemail + html7)
    html8 = """
     </p>
        </div>
        <hr style="clear:both;">
        <div class="clear"></div>

        <dl>
            <dd class="clear"></dd>

            <dt>Education</dt><dd>

    """
    htmlfile.write(gheadln + html8)

    if edu > 0:
        for ed in range(0, edu + 1):
            guniv = univ[ed].get()
            gschool = school[ed].get()
            ggpa = gpa[ed].get()
            gloc = loc[ed].get()
            gcourse = course[ed].get()
            gmajor = major[ed].get()
            gdate = datee[ed].get()
            html10 = """
            <h2>
            """
            html11 = ","
            html12 = "</h2><p><strong>"
            html13 = "</strong>"
            html14 = "<br /><strong>Major:</strong>"
            html15 = "</p>"
            htmlfile.write(
                html10 + guniv + html11 + gschool + html11 + gloc + html12 + gcourse + html13 + gdate + html14 + gmajor + html15)


    elif edu == 0:
        guniv = univ[0].get()
        gschool = school[0].get()
        ggpa = gpa[0].get()
        gloc = loc[0].get()
        gcourse = course[0].get()
        gmajor = major[0].get()
        gdate = datee[0].get()
        html10 = """
            <h2>
            """
        html11 = ","
        html12 = "</h2><p><strong>"
        html13 = "</strong>"
        html14 = "<br /><strong>Major:</strong>"
        html15 = "</p></dd>"
        htmlfile.write(html10)
        htmlfile.write(guniv)
        htmlfile.write(html11)
        htmlfile.write(gschool)
        htmlfile.write(html11)
        htmlfile.write(gloc)
        htmlfile.write(html12)
        htmlfile.write(gcourse)
        htmlfile.write(html13)
        htmlfile.write(gdate)
        htmlfile.write(html14)
        htmlfile.write(gmajor)
        htmlfile.write(html15)
    html16 = "</dd>"
    htmlfile.write(html16)
    html17 = """<dd class="clear"></dd>
    <dt>Experience</dt>
    <dd>
    """
    htmlfile.write(html17)

    if exp > 0:
        for ed in range(0, exp + 1):
            gcomp = company[ed].get()
            glocexp = locexp[ed].get()
            gposi = posi[ed].get()
            gsd = startdate[ed].get()
            ged = enddate[ed].get()
            gdetails = detailss[ed].get(1.0, END)
            html10 = "<h2>"
            html11 = "<br><span>"
            html12 = " - "
            html13 = """</span><br></h2>

                <ul>
                    <li>"""
            html14 = "</li></ul>"

            htmlfile.write(
                html10 + gcomp + html11 + gposi + html12 + glocexp + html12 + gsd + html12 + ged + html13 + gdetails + html14)


    elif exp == 0:
        gcomp = company[0].get()
        glocexp = locexp[0].get()
        gposi = posi[0].get()
        gsd = startdate[0].get()
        ged = enddate[0].get()
        gdetails = detailss[0].get(1.0, END)
        html10 = "<h2>"
        html11 = "<br><span>"
        html12 = " - "
        html13 = """</span><br></h2>
            <ul>
                    <li>"""
        html14 = "</li></ul>"

        htmlfile.write(
            html10 + gcomp + html11 + gposi + html12 + glocexp + html12 + gsd + html12 + ged + html13 + gdetails + html14)

    htmlfile.write(html16)
    html17 = """<dd class="clear"></dd>
    <dt>Skills</dt>
    <dd>
    """
    htmlfile.write(html17)

    if skil > 0:
        for sksk in range(0, skil + 1):
            gstitle = stitle[sksk].get()
            gsdetail = sdetail[sksk].get()
            html10 = "<h2>"
            html11 = "</h2><p>"
            html12 = "</p>"
            htmlfile.write(
                html10 + gstitle + html11 + gsdetail + html12)

    elif skil == 0:
        gstitle = stitle[0].get()
        gsdetail = sdetail[0].get()
        html10 = "<h2>"
        html11 = "</h2><p>"
        html12 = "</p>"
        htmlfile.write(
            html10 + gstitle + html11 + gsdetail + html12)

    htmlfile.write(html16)
    html17 = """<dd class="clear"></dd>
    <dt>Projects</dt>
    <dd>
    """
    htmlfile.write(html17)

    if projcount > 0:
        for pjpj in range(0, projcount + 1):
            gprojtitle = projname[pjpj].get()
            gprojdetail = projdetail[pjpj].get()
            html10 = "<h2>"
            html11 = "</h2><p>"
            html12 = "</p>"
            htmlfile.write(
                html10 + gprojtitle + html11 + gprojdetail + html12)

    elif projcount == 0:
        gprojtitle = projname[0].get()
        gprojdetail = projdetail[0].get()
        html10 = "<h2>"
        html11 = "</h2><p>"
        html12 = "</p>"
        htmlfile.write(
            html10 + gprojtitle + html11 + gprojdetail + html12)

    httml = """
        <dd class="clear"></dd>



                <dt>References</dt>
                <dd>Available on request</dd>




            </dl>
            <div class="clear"></div>
            <hr style="margin-bottom:0px;margin-top:25px">


        </div>
        <footer>
            <h1><a href="
        """
    htmlfile.write(httml + gwl2)
    httml = """
    ">
    """
    httml1 = "</h1><br></footer></body></html>"
    htmlfile.write(httml + gname + httml1)
    htmlfile.close()


    #word_file

    document = Document()

    # name
    head = document.add_paragraph()

    head.paragraph_format.space_before = Pt(0)
    head.paragraph_format.space_after = Pt(0)
    head.size = Pt(14)
    head.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    head.add_run(gname,0).bold=True


    #address and phone
    paragrapha = document.add_paragraph(gaddres + " · " + gphonee)

    paragrapha.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    paragrapha.paragraph_format.space_before = Pt(0)
    paragrapha.paragraph_format.space_after = Pt(0)
    paragrapha.size = Pt(12)

    #email and weblinks
    paragrapha = document.add_heading(gweblink + " · " + gemail,0)

    paragrapha.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    paragrapha.paragraph_format.space_before = Pt(0)
    paragrapha.paragraph_format.space_after = Pt(8)
    paragrapha.size = Pt(10)


    edunmaee = document.add_paragraph()
    edunmaee.size = Pt(12)

    edunmaee.paragraph_format.space_before = Pt(0)
    edunmaee.paragraph_format.space_after = Pt(0)
    edunmaee.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    edunmaee.underline= WD_UNDERLINE.SINGLE
    edunmaee.add_run('EDUCATION').bold = True



    if edu > 0:
        for ed in range(0, edu + 1):
            guniv = univ[ed].get()
            gschool = school[ed].get()
            ggpa = gpa[ed].get()
            gloc = loc[ed].get()
            gcourse = course[ed].get()
            gmajor = major[ed].get()
            gdate = datee[ed].get()
            paragrapha = document.add_paragraph()
            paragrapha.paragraph_format.space_before = Pt(0)
            paragrapha.paragraph_format.space_after = Pt(0)
            paragrapha.size = Pt(12)
            paragrapha.add_run(guniv + "," + gschool + ",").bold = True
            paragrapha.add_run(gloc)

            paragrapha = document.add_paragraph()
            paragrapha.paragraph_format.space_before = Pt(0)
            paragrapha.paragraph_format.space_after = Pt(0)
            paragrapha.size = Pt(12)
            paragrapha.add_run(gcourse + "," + gmajor + " | " + ggpa + "     " + gdate)

    elif edu == 0:
        guniv = univ[0].get()
        gschool = school[0].get()
        ggpa = gpa[0].get()
        gloc = loc[0].get()
        gcourse = course[0].get()
        gmajor = major[0].get()
        gdate = datee[0].get()
        paragrapha = document.add_paragraph()
        paragrapha.paragraph_format.space_before = Pt(0)
        paragrapha.paragraph_format.space_after = Pt(0)
        paragrapha.size = Pt(12)
        paragrapha.add_run(guniv + "," + gschool + ",").bold = True
        paragrapha.add_run(gloc)

        paragrapha = document.add_paragraph()
        paragrapha.paragraph_format.space_before = Pt(0)
        paragrapha.paragraph_format.space_after = Pt(0)
        paragrapha.size = Pt(12)
        paragrapha.add_run(gcourse + "," + gmajor + " | " + ggpa + "     " + gdate)





    edunmaee = document.add_paragraph()
    edunmaee.size = Pt(12)

    edunmaee.paragraph_format.space_before = Pt(0)
    edunmaee.paragraph_format.space_after = Pt(0)
    edunmaee.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    edunmaee.underline= WD_UNDERLINE.SINGLE
    edunmaee.add_run('EXPERIENCE').bold = True


    if exp > 0:
        for ed in range(0, exp + 1):
            gcomp = company[ed].get()
            glocexp = locexp[ed].get()
            gposi = posi[ed].get()
            gsd = startdate[ed].get()
            ged = enddate[ed].get()
            gdetails = detailss[ed].get(1.0, END)
            paragrapha = document.add_paragraph()
            paragrapha.paragraph_format.space_before = Pt(0)
            paragrapha.paragraph_format.space_after = Pt(0)
            paragrapha.size = Pt(12)
            paragrapha.add_run(gcomp + "," + glocexp).bold = True

            paragrapha = document.add_paragraph()
            paragrapha.paragraph_format.space_before = Pt(0)
            paragrapha.paragraph_format.space_after = Pt(0)
            paragrapha.size = Pt(12)
            paragrapha.underline = WD_UNDERLINE.SINGLE
            paragrapha.add_run(gposi)
            paragrapha.underline = WD_UNDERLINE.NONE
            paragrapha.add_run(", "+ gsd + " - " + ged)

            paragrapha = document.add_paragraph()
            paragrapha.paragraph_format.space_before = Pt(0)
            paragrapha.paragraph_format.space_after = Pt(0)
            paragrapha.size = Pt(12)
            paragrapha.add_run(gdetails)


    elif exp == 0:
            gcomp = company[0].get()
            glocexp = locexp[0].get()
            gposi = posi[0].get()
            gsd = startdate[0].get()
            ged = enddate[0].get()
            gdetails = detailss[0].get(1.0,END)
            paragrapha = document.add_paragraph()
            paragrapha.paragraph_format.space_before = Pt(0)
            paragrapha.paragraph_format.space_after = Pt(0)
            paragrapha.size = Pt(12)
            paragrapha.add_run(gcomp + "," + glocexp).bold = True

            paragrapha = document.add_paragraph()
            paragrapha.paragraph_format.space_before = Pt(0)
            paragrapha.paragraph_format.space_after = Pt(0)
            paragrapha.size = Pt(12)
            paragrapha.underline = WD_UNDERLINE.SINGLE
            paragrapha.add_run(gposi)
            paragrapha.underline = WD_UNDERLINE.NONE
            paragrapha.add_run(", "+ gsd + " - " + ged)
            paragrapha = document.add_paragraph()
            paragrapha.paragraph_format.space_before = Pt(0)
            paragrapha.paragraph_format.space_after = Pt(0)
            paragrapha.size = Pt(12)
            paragrapha.add_run(gdetails)



    edunmaee = document.add_paragraph()
    edunmaee.size = Pt(12)

    edunmaee.paragraph_format.space_before = Pt(0)
    edunmaee.paragraph_format.space_after = Pt(0)
    edunmaee.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    edunmaee.underline= WD_UNDERLINE.SINGLE
    edunmaee.add_run('SKILLS').bold = True
    if skil > 0:
        for sksk in range(0, skil + 1):
            gstitle = stitle[sksk].get()
            gsdetail = sdetail[sksk].get()
            paragrapha = document.add_paragraph()
            paragrapha.paragraph_format.space_before = Pt(0)
            paragrapha.paragraph_format.space_after = Pt(0)
            paragrapha.size = Pt(12)
            paragrapha.add_run(gstitle).bold = True
            paragrapha.add_run(" : " + gsdetail).bold = False


    elif skil == 0:
        gstitle = stitle[0].get()
        gsdetail = sdetail[0].get()
        paragrapha = document.add_paragraph()
        paragrapha.paragraph_format.space_before = Pt(0)
        paragrapha.paragraph_format.space_after = Pt(0)
        paragrapha.size = Pt(12)
        paragrapha.add_run(gstitle).bold = True
        paragrapha.add_run(" : " + gsdetail).bold = False


    edunmaee = document.add_paragraph()
    edunmaee.size = Pt(12)

    edunmaee.paragraph_format.space_before = Pt(0)
    edunmaee.paragraph_format.space_after = Pt(0)
    edunmaee.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    edunmaee.underline= WD_UNDERLINE.SINGLE
    edunmaee.add_run('PROJECT').bold = True
    if projcount > 0:
        for pjpj in range(0, projcount + 1):
            gprojtitle = projname[pjpj].get()
            gprojdetail = projdetail[pjpj].get()
            paragrapha = document.add_paragraph()
            paragrapha.paragraph_format.space_before = Pt(0)
            paragrapha.paragraph_format.space_after = Pt(0)
            paragrapha.size = Pt(12)
            paragrapha.add_run(gprojtitle).bold = True
            paragrapha.add_run(" : " + gprojdetail).bold = False


    elif projcount == 0:
        gprojtitle = projname[0].get()
        gprojdetail = projdetail[0].get()
        paragrapha = document.add_paragraph()
        paragrapha.paragraph_format.space_before = Pt(0)
        paragrapha.paragraph_format.space_after = Pt(0)
        paragrapha.size = Pt(12)
        paragrapha.add_run(gprojtitle).bold = True
        paragrapha.add_run(" : " + gprojdetail).bold = False







    document.save('demo.docx')
    complete1()
    pdffi = path.abspath('.\demo.docx')
    convert_to_pdf(pdffi)



def convert_to_pdf(doc):
    word = client.DispatchEx("Word.Application")
    new_name = doc.replace(".docx", r".pdf")
    worddoc = word.Documents.Open(doc)
    worddoc.SaveAs(new_name, FileFormat=17)
    worddoc.Close()

def complete1():
    tkMessageBox.showinfo("Viola!!","All Files Generated")

def completeword():
    tkMessageBox.showinfo("Viola!!","Word File Generated.")

def completehtml():
    tkMessageBox.showinfo("Viola!!","HTML File Generated.")

def completepdf():
    tkMessageBox.showinfo("Viola!!","PDF File Generated.")


def create_tact():
    window = tk.Toplevel(root)
    window.title("Contact")
    window.resizable(width=FALSE, height=FALSE)
    window.geometry("600x500")
    global fromaddrs, username, password, msg
    ct1 = tk.Label(window, text="Email (GMAIL)", font=("Helvetica", 20))
    ct1.pack()
    fromaddrs = tk.Entry(window, width="60")
    fromaddrs.pack()
    b2 = tk.Label(window,
                  text="__________________________________________________________________________________________",
                  font=("Helvetica", 10))
    b2.pack()
    ct2 = tk.Label(window, text="Password", font=("Helvetica", 20))
    ct2.pack()
    password = tk.Entry(window, show="*", width="60")
    password.pack()
    b2 = tk.Label(window,
                  text="__________________________________________________________________________________________",
                  font=("Helvetica", 10))
    b2.pack()
    ct3 = tk.Label(window, text="Message and your number for me", font=("Helvetica", 20))
    ct3.pack()
    msg = tk.Text(window, height="9", width="80")
    msg.pack()
    b2 = tk.Label(window,
                  text="__________________________________________________________________________________________\n\n",
                  font=("Helvetica", 10))
    b2.pack()

    saend = tk.Button(window, text="Send", command=seend, font=("Helvetica", 30), fg="blue", relief=RAISED)
    saend.place(relx=.25, rely=.8)
    clo = tk.Button(window, text="Close", command=window.destroy, font=("Helvetica", 30), fg="blue", relief=RAISED)
    clo.place(relx=.60, rely=.8)
    fut = Label(window, text="**Only GMAIL supported Currently")
    fut.pack(side=BOTTOM)


def seend():
    global fromaddrs, username, password, msg
    toa = 'dhruv@dhruvgandhi.me'
    user = fromaddrs.get()
    pas = password.get()
    mseg = msg.get(1.0, tk.END)
    # The actual mail send
    server = smtplib.SMTP('smtp.gmail.com:587')
    server.starttls()
    server.login(user, pas)
    server.sendmail(user, toa, mseg)
    server.quit()


def create_faq():
    window = tk.Toplevel(root)
    window.title("Frequently Asked Questions")
    window.geometry("800x600")

    scrollbar = Scrollbar(window)
    scrollbar.pack(side=RIGHT, fill=Y)

    text = Text(window, wrap=WORD, yscrollcommand=scrollbar.set)
    text.place(relwidth=.97, relheight=1)
    text.tag_configure('question', font=('Verdana', 15, 'bold'), justify=LEFT)
    text.tag_configure('answer', font=('Verdana', 10), justify=LEFT)
    text.tag_configure('big', font=('Verdana', 30, 'bold', 'underline'), justify=CENTER)
    text.insert(tk.END, '\nFrequenlty Asked Question\n', 'big')
    quote = """
    Q.1 What will I Present Here?
    """
    anss = """
    I will write basic tip to run the application. If error is face you can even send me mail and attachment of screenshots of error through your gmail. I promise your password is safe.
    """
    text.insert(tk.END, quote, 'question')
    text.insert(END, anss, 'answer')

    text.config(state=DISABLED)
    scrollbar.config(command=text.yview)


def create_faq1():
    window = tk.Toplevel(mai)
    window.title("Frequently Asked Questions")
    window.geometry("800x600")

    scrollbar = Scrollbar(window)
    scrollbar.pack(side=RIGHT, fill=Y)

    text = Text(window, wrap=WORD, yscrollcommand=scrollbar.set)
    text.place(relwidth=.97, relheight=1)
    text.tag_configure('question', font=('Verdana', 15, 'bold'), justify=LEFT)
    text.tag_configure('answer', font=('Verdana', 10), justify=LEFT)
    text.tag_configure('big', font=('Verdana', 30, 'bold', 'underline'), justify=CENTER)
    text.insert(tk.END, '\nFrequenlty Asked Question\n', 'big')
    quote = """
    Q.1 What ill I Present Here?
    """
    anss = """
    I will write basic tip to run the application. If error is face you can even send me mail and attachment of screenshots of error through your gmail. I promise your password is safe.
    """
    text.insert(tk.END, quote, 'question')
    text.insert(END, anss, 'answer')

    text.config(state=DISABLED)
    scrollbar.config(command=text.yview)


def create_err():
    window = tk.Toplevel(root)
    window.title("Report Error")
    window.resizable(width=FALSE, height=FALSE)
    window.geometry("600x500")
    global fromaddrs, username, password, filepath
    ct1 = tk.Label(window, text="Email", font=("Helvetica", 20))
    ct1.pack()
    fromaddrs = tk.Entry(window, width="60")
    fromaddrs.pack()
    b2 = tk.Label(window,
                  text="__________________________________________________________________________________________",
                  font=("Helvetica", 10))
    b2.pack()
    ct2 = tk.Label(window, text="Password", font=("Helvetica", 20))
    ct2.pack()
    password = tk.Entry(window, show="*", width="60")
    password.pack()
    b2 = tk.Label(window,
                  text="__________________________________________________________________________________________",
                  font=("Helvetica", 10))
    b2.pack()
    ct3 = tk.Label(window, text="Attach Screenshot Of Error", font=("Helvetica", 20))
    ct3.pack()
    msg = tk.Button(window, text="Browse", command=filee, width=400)
    msg.pack()
    fpn = Label(window, text="Selected File", font=("Helvetica", 25))
    fpn.pack()
    filepath = Text(window, width=50, height=1)
    filepath.pack()
    b2 = tk.Label(window,
                  text="__________________________________________________________________________________________\n\n",
                  font=("Helvetica", 10))
    b2.pack()

    saend = tk.Button(window, text="Send", command=snder, font=("Helvetica", 30), fg="blue", relief=RAISED)
    saend.place(relx=.25, rely=.8)
    clo = tk.Button(window, text="Close", command=window.destroy, font=("Helvetica", 30), fg="blue", relief=RAISED)
    clo.place(relx=.60, rely=.8)
    fut = Label(window, text="**Only GMAIL supported Currently")
    fut.pack(side=BOTTOM)
    window.mainloop()


def filee():
    global fname, filepath
    fname = askopenfilename(filetypes=(("Text files", "*.txt"), ("JPEG Giles", "*.jpg,*.jpeg, *.png"),
                                       ("HTML files", "*.html;*.htm"),
                                       ("All files", "*.*")))
    filepath.insert(tk.END, fname)
    filepath.config(state=DISABLED)


def snder():
    global fromaddrs, username, password, msg
    toa = 'dhruv@dhruvgandhi.me'
    user = fromaddrs.get()
    pas = password.get()
    with open(fname, "rb") as image_file:
        encoded_string = base64.b64encode(image_file.read())
    server = smtplib.SMTP('smtp.gmail.com:587')
    server.starttls()
    server.login(user, pas)
    server.sendmail(user, toa, encoded_string)
    server.quit()


def root_app():
    mai.destroy()
    root = tk.Tk()
    root.title("Resume Builder")
    root.geometry("1024x768")
    root.resizable(width=FALSE, height=FALSE)
    background_image = PhotoImage(file=r"E:/aas.gif")
    background_label = Label(root, image=background_image)
    background_label.place(x=0, y=0, relwidth=1, relheight=1)
    l = tk.Label(root, text="RESUME BUILDER", font=("Helvetica", 50), fg="red", padx=5, pady=5)
    l.pack()
    details = Text(root, wrap=WORD, height=5)
    details.tag_configure('big', font=('Verdana'), justify=CENTER)
    resdetails = """
    Tired Of Creating Resume On Your Own \n Want a quick Solution? \n Try this app and please leave your feedback..
        """
    details.insert(END, resdetails, 'big')
    details.place(relx=.18, rely=.29)
    details.config(state=DISABLED)
    dev = tk.Label(root, text=" Developed By \nDhruvil Gandhi \nJacky Patel \nShraddha Thakker", font=("Helvetica", 20),
                   justify=CENTER)
    dev.place(relx=.38, rely=.5)

    star = tk.Button(root, text="Start", command=create_star, font=("Helvetica", 25), fg="blue")
    star.place(relx=.07, rely=.7)
    err = tk.Button(root, text="Report  Error", command=create_err, font=("Helvetica", 25), fg="blue")
    err.place(relx=.26, rely=.7)
    tact = tk.Button(root, text="Contact", command=create_tact, font=("Helvetica", 25), fg="blue")
    tact.place(relx=.58, rely=.7)
    faq = tk.Button(root, text="FAQ", command=create_faq, font=("Helvetica", 25), fg="blue")
    faq.place(relx=.80, rely=.7)

    root.mainloop()


global toa
root = tk.Tk()
root.title("Resume Builder")
root.geometry("1024x768")
root.resizable(width=FALSE, height=FALSE)
background_image = PhotoImage(file=r"E:/aas.gif")
background_label = Label(root, image=background_image)
background_label.place(x=0, y=0, relwidth=1, relheight=1)
l = tk.Label(root, text="RESUME BUILDER", font=("Helvetica", 50), fg="red", padx=5, pady=5)
l.pack()
details = Text(root, wrap=WORD, height=5)
details.tag_configure('big', font=('Verdana'), justify=CENTER)
resdetails = """
Tired Of Creating Resume On Your Own \n Want a quick Solution? \n Try this app and please leave your feedback..
    """
details.insert(END, resdetails, 'big')
details.place(relx=.18, rely=.29)
details.config(state=DISABLED)
dev = tk.Label(root, text=" Developed By \nDhruvil Gandhi \nJacky Patel \nShraddha Thakker", font=("Helvetica", 20),
               justify=CENTER)
dev.place(relx=.38, rely=.5)

star = tk.Button(root, text="Start", command=create_star, font=("Helvetica", 25), fg="blue")
star.place(relx=.07, rely=.7)
err = tk.Button(root, text="Report  Error", command=create_err, font=("Helvetica", 25), fg="blue")
err.place(relx=.26, rely=.7)
tact = tk.Button(root, text="Contact", command=create_tact, font=("Helvetica", 25), fg="blue")
tact.place(relx=.58, rely=.7)
faq = tk.Button(root, text="FAQ", command=create_faq, font=("Helvetica", 25), fg="blue")
faq.place(relx=.80, rely=.7)

root.mainloop()